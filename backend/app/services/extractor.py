# Document extractor service for reading PDF, Word, etc.
import os
import tempfile
from typing import Dict, Any, Optional
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
# TODO: install python-pptx for PowerPoint support
# from pptx import Presentation
import logging

logger = logging.getLogger(__name__)

class DocumentExtractor:
    """Extracts text content from various document formats"""
    
    SUPPORTED_FORMATS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        '.ppt': 'application/vnd.ms-powerpoint',
        '.txt': 'text/plain'
    }
    
    @classmethod
    def is_supported_format(cls, filename: str) -> bool:
        """Check if file format is supported"""
        file_ext = Path(filename).suffix.lower()
        return file_ext in cls.SUPPORTED_FORMATS
    
    @classmethod
    def extract_text(cls, file_path: str, filename: str) -> Dict[str, Any]:
        """Extract text content from document"""
        try:
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.pdf':
                return cls._extract_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return cls._extract_word(file_path)
            elif file_ext in ['.pptx', '.ppt']:
                return cls._extract_powerpoint(file_path)
            elif file_ext == '.txt':
                return cls._extract_text_file(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise
    
    @staticmethod
    def _extract_pdf(file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file using PyMuPDF for better performance and bounding box support"""
        text_content = []
        metadata = {}
        
        try:
            # Open PDF with PyMuPDF
            pdf_document = fitz.open(file_path)
            
            # Extract metadata
            pdf_metadata = pdf_document.metadata
            if pdf_metadata:
                metadata = {
                    'title': pdf_metadata.get('title', ''),
                    'author': pdf_metadata.get('author', ''),
                    'subject': pdf_metadata.get('subject', ''),
                    'creator': pdf_metadata.get('creator', ''),
                    'producer': pdf_metadata.get('producer', ''),
                    'pages': pdf_document.page_count
                }
            
            # Extract text from each page with bounding box information
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                
                # Get text blocks with bounding boxes for highlighting support
                text_blocks = page.get_text("dict")
                
                page_text_content = []
                full_page_text = ""
                
                # Process each text block
                for block in text_blocks["blocks"]:
                    if "lines" in block:  # Text block
                        for line in block["lines"]:
                            for span in line["spans"]:
                                if span["text"].strip():
                                    # Store text with bounding box for highlighting
                                    text_item = {
                                        'text': span["text"],
                                        'bbox': span["bbox"],  # [x0, y0, x1, y1]
                                        'font': span["font"],
                                        'size': span["size"],
                                        'flags': span["flags"]
                                    }
                                    page_text_content.append(text_item)
                                    full_page_text += span["text"] + " "
                
                if full_page_text.strip():
                    text_content.append({
                        'page': page_num + 1,
                        'content': full_page_text.strip(),
                        'text_blocks': page_text_content,  # For highlighting support
                        'page_bbox': page.rect  # Page dimensions
                    })
            
            pdf_document.close()
                
        except Exception as e:
            logger.error(f"Error reading PDF file: {str(e)}")
            raise
        
        return {
            'content': text_content,
            'metadata': metadata,
            'total_pages': len(text_content),
            'format': 'pdf'
        }
    
    @staticmethod
    def _extract_word(file_path: str) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'creator': doc.core_properties.creator or '',
                'revision': doc.core_properties.revision or 0,
                'paragraphs': len(doc.paragraphs)
            }
            
            # Extract text content
            text_content = []
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append({
                        'paragraph': para_num + 1,
                        'content': paragraph.text.strip(),
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
            
            return {
                'content': text_content,
                'metadata': metadata,
                'total_paragraphs': len(text_content),
                'format': 'word'
            }
            
        except Exception as e:
            logger.error(f"Error reading Word document: {str(e)}")
            raise
    
    @staticmethod
    def _extract_powerpoint(file_path: str) -> Dict[str, Any]:
        """Extract text from PowerPoint presentation"""
        try:
            # For now, return a placeholder since python-pptx is not installed
            # TODO: Install python-pptx and implement proper extraction
            logger.warning("PowerPoint extraction not fully implemented - python-pptx not installed")
            
            return {
                'content': [{'slide': 1, 'content': 'PowerPoint extraction not implemented yet'}],
                'metadata': {'slides': 1, 'note': 'python-pptx not installed'},
                'total_slides': 1,
                'format': 'powerpoint'
            }
            
        except Exception as e:
            logger.error(f"Error reading PowerPoint presentation: {str(e)}")
            raise
    
    @staticmethod
    def _extract_text_file(file_path: str) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Split into paragraphs
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            
            text_content = []
            for para_num, paragraph in enumerate(paragraphs):
                text_content.append({
                    'paragraph': para_num + 1,
                    'content': paragraph
                })
            
            return {
                'content': text_content,
                'metadata': {
                    'encoding': 'utf-8',
                    'paragraphs': len(text_content)
                },
                'total_paragraphs': len(text_content),
                'format': 'text'
            }
            
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            raise
